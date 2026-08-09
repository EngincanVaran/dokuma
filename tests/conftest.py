from __future__ import annotations

from email.message import EmailMessage
from pathlib import Path

import docx
import openpyxl
import pytest
import xlwt
from fpdf import FPDF
from PIL import Image


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """A real, valid, text-based 3-page PDF - not a hand-rolled byte string,
    so page count / text-layer detection are exercised against genuine PDF
    structure."""
    pdf = FPDF()
    for i in range(3):
        pdf.add_page()
        pdf.set_font("Helvetica", size=14)
        pdf.cell(text=f"Page {i + 1} of the dokuma test fixture.")

    path = tmp_path / "sample.pdf"
    pdf.output(str(path))
    return path


@pytest.fixture
def dense_text_pdf(tmp_path: Path) -> Path:
    """A PDF with a real paragraph of text, not one short line per page.

    Separate from `sample_pdf` on purpose: pdf-inspector's `process_pdf`
    (used by the extraction engine, unlike `inspect()`'s `detect_pdf`) is
    sensitive to per-page text density - confirmed by direct testing,
    `sample_pdf`'s one-short-line-per-page content gets flagged as
    `pages_needing_ocr` and extracts to empty markdown despite pdfplumber
    reading it back just fine, while this fixture's denser paragraph
    extracts cleanly at confidence 1.0. Real behavior, not a bug to work
    around - this fixture exists to test the *working* path; the sparse
    path is exercised directly against `sample_pdf` in
    test_engine_pdf_inspector.py.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 8, "This is a real paragraph of test content. " * 8)

    path = tmp_path / "dense.pdf"
    pdf.output(str(path))
    return path


@pytest.fixture
def table_pdf(tmp_path: Path) -> Path:
    """A PDF with a real paragraph plus a real table, for engines (like
    pdfplumber) that detect table structure rather than just reading a
    pre-existing markdown pipe-table."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 8, "Intro paragraph before the table.")
    with pdf.table() as table:
        row = table.row()
        row.cell("Name")
        row.cell("Value")
        row = table.row()
        row.cell("widgets")
        row.cell("42")

    path = tmp_path / "table.pdf"
    pdf.output(str(path))
    return path


@pytest.fixture
def text_table_text_pdf(tmp_path: Path) -> Path:
    """Text both before AND after the table - table_pdf alone can't catch a
    "text region always comes after tables" bug, since it has no closing
    text to test against."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 8, "Intro paragraph before the table.")
    with pdf.table() as table:
        row = table.row()
        row.cell("Name")
        row.cell("Value")
        row = table.row()
        row.cell("widgets")
        row.cell("42")
    pdf.multi_cell(0, 8, "Closing paragraph after the table.")

    path = tmp_path / "text_table_text.pdf"
    pdf.output(str(path))
    return path


@pytest.fixture
def image_pdf(tmp_path: Path) -> Path:
    """Text, then a real embedded picture, then more text - for
    PdfPlumberEngine's image-region extraction. The picture is a small
    solid-color PIL image, not a photo - content doesn't matter, only that
    it's a genuine raster XObject pdfplumber's `page.images` will detect,
    not a hand-rolled byte string."""
    picture = Image.new("RGB", (60, 40), color=(200, 30, 30))

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 8, "Intro paragraph before the picture.")
    pdf.ln(5)
    pdf.image(picture, x=20, w=40)
    pdf.ln(5)
    pdf.multi_cell(0, 8, "Closing paragraph after the picture.")

    path = tmp_path / "image.pdf"
    pdf.output(str(path))
    return path


@pytest.fixture
def table_and_image_pdf(tmp_path: Path) -> Path:
    """Text, a table, an image, then more text - so real reading order can
    be verified across all three block kinds PdfPlumberEngine merges by
    vertical position, not just text+table."""
    picture = Image.new("RGB", (60, 40), color=(30, 120, 200))

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 8, "Intro paragraph.")
    with pdf.table() as table:
        row = table.row()
        row.cell("Name")
        row.cell("Value")
        row = table.row()
        row.cell("widgets")
        row.cell("42")
    pdf.ln(5)
    pdf.image(picture, x=20, w=40)
    pdf.ln(45)
    pdf.multi_cell(0, 8, "Closing paragraph.")

    path = tmp_path / "table_and_image.pdf"
    pdf.output(str(path))
    return path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """python-docx doesn't populate Word's cached page-count field (only
    Word itself does, on save) - so this fixture's page_count is expected
    to come back None. That's exercised explicitly in the tests, not just
    incidental."""
    document = docx.Document()
    document.core_properties.title = "Dokuma Test Document"
    document.core_properties.author = "dokuma"
    document.add_heading("Introduction", level=1)
    document.add_paragraph("This is the first paragraph of the test fixture.")
    document.add_heading("Details", level=2)
    document.add_paragraph("This is the second paragraph, with more words in it.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"

    path = tmp_path / "sample.docx"
    document.save(str(path))
    return path


@pytest.fixture
def sample_xlsx(tmp_path: Path) -> Path:
    workbook = openpyxl.Workbook()
    workbook.properties.title = "Dokuma Test Workbook"
    workbook.properties.creator = "dokuma"

    sheet1 = workbook.active
    sheet1.title = "Data"
    sheet1["A1"] = "Name"
    sheet1["B1"] = "Value"
    sheet1["A2"] = "widgets"
    sheet1["B2"] = 42

    sheet2 = workbook.create_sheet("Notes")
    sheet2["A1"] = "just a note"

    path = tmp_path / "sample.xlsx"
    workbook.save(str(path))
    return path


@pytest.fixture
def sample_html(tmp_path: Path) -> Path:
    html = """<!DOCTYPE html>
<html>
<head><title>Dokuma Test Page</title></head>
<body>
<h1>Welcome</h1>
<p>This is a paragraph with some words in it.</p>
<h2>Section Two</h2>
<p>Another paragraph here, also with a few words.</p>
</body>
</html>"""
    path = tmp_path / "sample.html"
    path.write_text(html, encoding="utf-8")
    return path


@pytest.fixture
def html_with_table(tmp_path: Path) -> Path:
    html = """<!DOCTYPE html>
<html>
<head><title>Report</title></head>
<body>
<h1>Welcome</h1>
<p>Intro paragraph.</p>
<table>
<tr><th>Name</th><th>Value</th></tr>
<tr><td>widgets</td><td>42</td></tr>
</table>
<h2>Section Two</h2>
<p>Closing paragraph.</p>
</body>
</html>"""
    path = tmp_path / "with_table.html"
    path.write_text(html, encoding="utf-8")
    return path


@pytest.fixture
def sample_markdown(tmp_path: Path) -> Path:
    text = """# Dokuma Test Doc

Some intro text with a handful of words.

## Section Two

More words go here, in this second section.

### Subsection

Even more text.
"""
    path = tmp_path / "sample.md"
    path.write_text(text, encoding="utf-8")
    return path


@pytest.fixture
def markdown_with_table(tmp_path: Path) -> Path:
    text = """# Welcome

Intro paragraph.

| Name | Value |
|---|---|
| widgets | 42 |

## Section Two

Closing paragraph.
"""
    path = tmp_path / "with_table.md"
    path.write_text(text, encoding="utf-8")
    return path


@pytest.fixture
def sample_csv(tmp_path: Path) -> Path:
    text = "name,value,note\nwidgets,42,in stock\ngadgets,7,backordered\n"
    path = tmp_path / "sample.csv"
    path.write_text(text, encoding="utf-8")
    return path


@pytest.fixture
def sample_tsv(tmp_path: Path) -> Path:
    text = "name\tvalue\tnote\nwidgets\t42\tin stock\ngadgets\t7\tbackordered\n"
    path = tmp_path / "sample.tsv"
    path.write_text(text, encoding="utf-8")
    return path


@pytest.fixture
def sample_xls(tmp_path: Path) -> Path:
    workbook = xlwt.Workbook()
    sheet1 = workbook.add_sheet("Data")
    sheet1.write(0, 0, "Name")
    sheet1.write(0, 1, "Value")
    sheet1.write(1, 0, "widgets")
    sheet1.write(1, 1, 42)
    workbook.add_sheet("Notes")

    path = tmp_path / "sample.xls"
    workbook.save(str(path))
    return path


@pytest.fixture
def sample_eml(tmp_path: Path) -> Path:
    message = EmailMessage()
    message["From"] = "sender@example.com"
    message["To"] = "recipient@example.com"
    message["Subject"] = "Dokuma Test Email"
    message.set_content("This is the body of the test email, with a few words in it.")
    message.add_attachment(
        b"fake pdf bytes", maintype="application", subtype="pdf", filename="report.pdf"
    )

    path = tmp_path / "sample.eml"
    path.write_bytes(bytes(message))
    return path
