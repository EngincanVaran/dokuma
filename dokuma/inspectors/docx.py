"""DOCX inspection via python-docx, plus Word's cached page count.

DOCX has no true page count in the file structure - pages are a render-time
concept (depend on page size, margins, fonts). Real Microsoft Word caches
its last-computed page count in docProps/app.xml when the file is saved;
python-docx doesn't expose that field, so it's read directly (a .docx is a
zip archive).

Important caveat, confirmed by testing against python-docx's own output:
any docx NOT actually saved by real Word - including ones python-docx
itself creates - carries a stale, template-inherited Pages value (often
literally "1", alongside other clearly-boilerplate fields like
Application="Microsoft Macintosh Word" and Words="0" that don't match the
real content at all). There is no reliable way to tell "genuinely 1 page"
apart from "never touched by Word, template default". Treat this field as
"Word's last real save, if any" - not as ground truth for files from
python-docx, LibreOffice, Google Docs exports, etc.
"""

from __future__ import annotations

import zipfile
from pathlib import Path
from xml.etree import ElementTree

from dokuma.types import DocumentInfo

_APP_XML_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}"


def _cached_page_count(path: Path) -> int | None:
    try:
        with zipfile.ZipFile(path) as zf, zf.open("docProps/app.xml") as f:
            root = ElementTree.parse(f).getroot()
            pages = root.find(f"{_APP_XML_NS}Pages")
            if pages is not None and pages.text:
                return int(pages.text)
    except (KeyError, zipfile.BadZipFile, ElementTree.ParseError, ValueError):
        pass
    return None


def inspect_docx(path: Path) -> DocumentInfo:
    import docx

    document = docx.Document(str(path))

    word_count = sum(len(p.text.split()) for p in document.paragraphs)
    # Default Word templates name heading styles "Heading 1".."Heading 9" -
    # a heuristic, not guaranteed for custom/localized templates.
    heading_count = sum(
        1 for p in document.paragraphs if p.style is not None and p.style.name.startswith("Heading")
    )

    props = document.core_properties
    metadata = {
        k: str(v)
        for k, v in {
            "author": props.author,
            "subject": props.subject,
            "created": props.created,
            "modified": props.modified,
        }.items()
        if v
    }

    return DocumentInfo(
        path=path,
        format="docx",
        size_bytes=path.stat().st_size,
        page_count=_cached_page_count(path),
        paragraph_count=len(document.paragraphs),
        table_count=len(document.tables),
        heading_count=heading_count,
        word_count=word_count,
        title=props.title or None,
        metadata=metadata,
    )
