"""DOCX extraction via python-docx.

Walks the document body directly (`document.element.body.iterchildren()`)
rather than using python-docx's `.paragraphs`/`.tables`, which are two
separate flat lists - not interleaved in real document order. Verified
directly: the body-walk approach correctly reproduces heading/paragraph/
table order; `.paragraphs`/`.tables` alone would silently lose it.

No bbox/page: DOCX is a flow document, not fixed-position - there's no
reliable per-paragraph page number without actually rendering the document
(Word computes pages at render/print time, same caveat as
`inspectors/docx.py`'s cached page count). `bbox`/`page` stay `None` on
every region, same honesty as PdfInspectorEngine.

Consecutive paragraphs are batched into one "text" region (split at
headings/tables), matching the PDF engines' "one region per contiguous
chunk" pattern rather than one region per paragraph. Headings get their own
"heading" category - real, free structure DOCX gives us that PDF text
extraction can't reliably infer.

Image regions: an inline picture lives in its own otherwise-empty
paragraph as a `w:drawing` element; found via `paragraph._element` (a
private python-docx attribute - there's no public accessor for a
paragraph's raw XML, same class of tradeoff as XlsxEngine reaching into
openpyxl's private `_images`) and resolved to real bytes through the
paragraph's relationship part (`paragraph.part.related_parts[rId]`), not
rendered - unlike PdfPlumberEngine, DOCX images are already stored as
whatever their original file format was, so `image_format` varies (PNG,
JPEG, ...) instead of always being PNG. Real bug fixed alongside this:
an image-only paragraph has `.text == ""`, so the pre-existing
`if not text: continue` would skip it *and* its image entirely - found
via real-document testing (confirmed empirically: zero trace of an
embedded picture anywhere in extraction output). Images are now found
before that skip check runs.

`ExtractConfig(extract_images=False)` skips `_paragraph_images()` (and
its relationship-part lookups) entirely, not just the resulting regions.
"""

from __future__ import annotations

from collections.abc import Iterator
from pathlib import Path
from typing import TYPE_CHECKING

from dokuma.config import ExtractConfig
from dokuma.engines.base import Engine
from dokuma.types import ExtractionResult, Region

if TYPE_CHECKING:
    from docx.document import Document
    from docx.text.paragraph import Paragraph as ParagraphType


def _paragraph_images(paragraph: ParagraphType) -> Iterator[tuple[bytes, str]]:
    from docx.oxml.ns import qn

    for blip in paragraph._element.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid is None:
            continue
        part = paragraph.part.related_parts[rid]
        yield part.blob, part.content_type.rsplit("/", 1)[-1]


def _table_to_html(table: object) -> str:
    from docx.table import Table

    assert isinstance(table, Table)
    body = []
    for row in table.rows:
        cells = "".join(f"<td>{cell.text}</td>" for cell in row.cells)
        body.append(f"<tr>{cells}</tr>")
    return "<table>" + "".join(body) + "</table>"


def _heading_level(style_name: str) -> int | None:
    """ "Heading 1" -> 1, "Heading 2" -> 2, ... "Heading" (no number, some
    templates) -> None rather than guessing."""
    suffix = style_name.removeprefix("Heading").strip()
    return int(suffix) if suffix.isdigit() else None


def _iter_body_elements(document: Document) -> Iterator[object]:
    """Yields Paragraph/Table objects in true document order."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


class DocxEngine(Engine):
    name = "python-docx"
    format = "docx"

    def _extract(self, path: Path, config: ExtractConfig) -> ExtractionResult:
        import docx
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        document = docx.Document(str(path))

        regions: list[Region] = []
        order = 0
        text_buffer: list[str] = []

        def flush_text() -> None:
            nonlocal order
            text = "\n".join(text_buffer).strip()
            if text:
                regions.append(Region(category="text", content=text, order=order))
                order += 1
            text_buffer.clear()

        for element in _iter_body_elements(document):
            if isinstance(element, Paragraph):
                if config.extract_images:
                    for image_bytes, image_format in _paragraph_images(element):
                        flush_text()
                        regions.append(
                            Region(
                                category="image",
                                content="",
                                order=order,
                                image_bytes=image_bytes,
                                image_format=image_format,
                            )
                        )
                        order += 1

                text = element.text.strip()
                if not text:
                    continue
                style_name = element.style.name if element.style is not None else ""
                if style_name.startswith("Heading"):
                    flush_text()
                    regions.append(
                        Region(
                            category="heading",
                            content=text,
                            order=order,
                            level=_heading_level(style_name),
                        )
                    )
                    order += 1
                else:
                    text_buffer.append(text)
            elif isinstance(element, Table):
                flush_text()
                regions.append(
                    Region(category="table", content=_table_to_html(element), order=order)
                )
                order += 1

        flush_text()

        return ExtractionResult(
            path=path,
            format=self.format,
            engine=self.name,
            regions=regions,
        )
