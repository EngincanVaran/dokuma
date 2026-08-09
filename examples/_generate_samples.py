"""Generates small, realistic sample documents for the scripts in this
folder. Not part of the public library - keeps every example
self-contained (no external sample files to hunt down or commit as
binary blobs) using the same generator libraries dokuma's own test suite
already relies on (fpdf2, python-docx, openpyxl, Pillow). Needs
`dokuma[all]` plus these dev-only generator libraries - see
examples/README.md.
"""

from __future__ import annotations

import io
from pathlib import Path


def _sample_pdf_text_and_table(pdf) -> None:
    pdf.add_page()
    pdf.set_font("Helvetica", size=16)
    pdf.cell(text="Quarterly Report")
    pdf.ln(14)
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(
        0,
        8,
        "Revenue grew steadily this quarter across every region, driven by "
        "strong renewal rates and a healthy pipeline of new accounts. "
        "The team expects this trend to continue into next quarter.",
    )
    pdf.ln(4)
    with pdf.table() as table:
        row = table.row()
        row.cell("Region")
        row.cell("Revenue")
        row = table.row()
        row.cell("EMEA")
        row.cell("120,000")
        row = table.row()
        row.cell("APAC")
        row.cell("98,000")


def make_sample_pdf(path: Path) -> None:
    """Text + a table, no embedded image - dense enough real text that
    pdf-inspector's fast classifier reads it as text_based, not
    needing OCR. Kept image-free on purpose: 04_images.py has its own
    generator so the "needs OCR" signal (sensitive to page image content,
    not just text density) doesn't leak into every other example."""
    from fpdf import FPDF

    pdf = FPDF()
    _sample_pdf_text_and_table(pdf)
    pdf.output(str(path))


def make_sample_pdf_with_image(path: Path) -> None:
    from fpdf import FPDF
    from PIL import Image

    picture = Image.new("RGB", (120, 80), color=(200, 30, 30))

    pdf = FPDF()
    _sample_pdf_text_and_table(pdf)
    pdf.ln(8)
    pdf.image(picture, x=20, w=50)
    pdf.output(str(path))


def make_sample_docx(path: Path) -> None:
    import docx
    from PIL import Image

    picture = io.BytesIO()
    Image.new("RGB", (120, 80), color=(30, 120, 200)).save(picture, format="PNG")
    picture.seek(0)

    document = docx.Document()
    document.add_heading("Project Handbook", level=1)
    document.add_paragraph("This handbook covers onboarding and team conventions.")
    document.add_heading("Team Roster", level=2)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Role"
    table.cell(1, 0).text = "Ada"
    table.cell(1, 1).text = "Engineer"
    document.add_picture(picture, width=docx.shared.Inches(1.5))
    document.save(str(path))


def make_sample_xlsx(path: Path) -> None:
    import openpyxl
    from openpyxl.drawing.image import Image as XlsxImage
    from PIL import Image

    picture = io.BytesIO()
    Image.new("RGB", (120, 80), color=(30, 180, 90)).save(picture, format="PNG")
    picture.seek(0)

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Inventory"
    sheet.append(["Item", "Quantity", "Price"])
    sheet.append(["Widgets", 42, 3.5])
    sheet.append(["Gadgets", 7, 12.0])
    image = XlsxImage(picture)
    image.anchor = "E2"
    sheet.add_image(image)
    workbook.save(str(path))
